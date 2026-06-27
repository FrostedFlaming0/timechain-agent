"""
docx_writer — markdown → .docx bytes for the write gate's generated formats.

The model authors MARKDOWN (its native register); this module produces the
binary Word document at proposal time, so `write_file` on a `.docx` path
seals a real document an attorney can open — never markdown-in-a-renamed-
file, never model-emitted XML.

Two backends, same contract as extractors.py in reverse:
  - python-docx when installed (proper Heading/List styles),
  - a minimal stdlib WordprocessingML writer otherwise (zipfile + XML with
    direct formatting) — plainer, but a fully valid document, so the
    feature works with zero new required dependencies.

Markdown coverage is deliberately the LOI/letter subset: #–###### headings,
paragraphs (consecutive lines joined, blank-line separated), **bold**,
*italic*, `-`/`*` bullets, `1.` numbered lists, and `---` rules. Tables and
images are out of scope until someone needs them.

Round-trip: extractors._extract_docx reads back what either backend writes.
"""

from __future__ import annotations

import io
import re
import zipfile
from xml.sax.saxutils import escape

# Heading sizes for the stdlib backend, in half-points (Word's w:sz unit).
_HEADING_SZ = {1: 36, 2: 30, 3: 26, 4: 24, 5: 22, 6: 22}

_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

_BULLET = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^(\d+)[.)]\s+(.*)$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def _inline_runs(text: str) -> list[tuple[str, bool, bool]]:
    """Split markdown inline emphasis into (text, bold, italic) runs."""
    runs = []
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            runs.append((part[2:-2], True, False))
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            runs.append((part[1:-1], False, True))
        else:
            runs.append((part, False, False))
    return runs


def _parse_blocks(md: str) -> list[dict]:
    """Line-based markdown parse into renderable blocks:
    {kind: heading|para|bullet|number|rule, text, level}."""
    blocks: list[dict] = []
    para: list[str] = []

    def flush_para():
        if para:
            blocks.append({"kind": "para", "text": " ".join(para)})
            para.clear()

    for raw in md.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_para()
            continue
        m = _HEADING.match(stripped)
        if m:
            flush_para()
            blocks.append({"kind": "heading", "level": len(m.group(1)),
                           "text": m.group(2).strip()})
            continue
        if stripped in ("---", "***", "___"):
            flush_para()
            blocks.append({"kind": "rule"})
            continue
        m = _BULLET.match(stripped)
        if m:
            flush_para()
            blocks.append({"kind": "bullet", "text": m.group(1)})
            continue
        m = _NUMBERED.match(stripped)
        if m:
            flush_para()
            blocks.append({"kind": "number", "text": m.group(2),
                           "n": m.group(1)})
            continue
        para.append(stripped)
    flush_para()
    return blocks


# --------------------------------------------------------- stdlib backend

_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _run_xml(text: str, bold: bool, italic: bool, sz: int = 0) -> str:
    props = ""
    if bold or italic or sz:
        props = ("<w:rPr>" + ("<w:b/>" if bold else "")
                 + ("<w:i/>" if italic else "")
                 + (f'<w:sz w:val="{sz}"/>' if sz else "") + "</w:rPr>")
    return (f'<w:r>{props}<w:t xml:space="preserve">'
            f"{escape(text)}</w:t></w:r>")


def _para_xml(runs: list[tuple[str, bool, bool]], sz: int = 0,
              all_bold: bool = False) -> str:
    body = "".join(_run_xml(t, b or all_bold, i, sz) for t, b, i in runs)
    return f"<w:p>{body}</w:p>"


def _stdlib_docx(md: str) -> bytes:
    parts = []
    for blk in _parse_blocks(md):
        if blk["kind"] == "heading":
            sz = _HEADING_SZ[blk["level"]]
            parts.append(_para_xml(_inline_runs(blk["text"]), sz=sz,
                                   all_bold=True))
        elif blk["kind"] == "rule":
            parts.append("<w:p/>")
        elif blk["kind"] == "bullet":
            parts.append(_para_xml([("• ", False, False)]
                                   + _inline_runs(blk["text"])))
        elif blk["kind"] == "number":
            parts.append(_para_xml([(f"{blk['n']}. ", False, False)]
                                   + _inline_runs(blk["text"])))
        else:
            parts.append(_para_xml(_inline_runs(blk["text"])))
    if not parts:
        parts.append("<w:p/>")
    document = (f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:document xmlns:w="{_W}"><w:body>'
                + "".join(parts) + "<w:sectPr/></w:body></w:document>")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _CONTENT_TYPES)
        z.writestr("_rels/.rels", _RELS)
        z.writestr("word/document.xml", document)
    return buf.getvalue()


# ------------------------------------------------------ python-docx backend

def _pythondocx_docx(md: str) -> bytes:
    import docx
    doc = docx.Document()
    for blk in _parse_blocks(md):
        if blk["kind"] == "heading":
            doc.add_heading(re.sub(r"\*+", "", blk["text"]),
                            level=min(blk["level"], 6))
            continue
        if blk["kind"] == "rule":
            doc.add_paragraph("")
            continue
        style = {"bullet": "List Bullet", "number": "List Number"}.get(
            blk["kind"])
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:                     # template lacks the list style
            p = doc.add_paragraph()
        for text, bold, italic in _inline_runs(blk["text"]):
            run = p.add_run(text)
            run.bold = bold or None
            run.italic = italic or None
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------- public API

def markdown_to_docx(md: str) -> tuple[bytes, str]:
    """Convert markdown to .docx bytes. Returns (bytes, backend) where
    backend is "python-docx" or "stdlib". Never raises for missing
    optional libraries — the stdlib writer always works."""
    try:
        return _pythondocx_docx(md), "python-docx"
    except ImportError:
        return _stdlib_docx(md), "stdlib"


# Formats write_file can generate, keyed by target extension. xlsx etc. can
# join this map later without touching the write gate again.
GENERATED_FORMATS = {".docx": markdown_to_docx}
